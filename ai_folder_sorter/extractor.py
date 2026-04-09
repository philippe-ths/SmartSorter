"""
File content extraction for SmartSorter.

This module provides a unified interface for extracting text content from
various file formats. Each extractor follows a Protocol interface, making
it easy to add support for new file types.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Protocol, Tuple, runtime_checkable


@dataclass(frozen=True)
class ExtractedContent:
    """
    Result of extracting content from a file.
    
    Attributes:
        text: The extracted text content
        method: The extraction method used (e.g., "pdf-text", "docx-text")
        truncated: Whether the content was truncated due to size limits
        is_full_content: Whether this represents the complete file content
        metadata: Optional additional metadata from extraction
    """
    text: str
    method: str
    truncated: bool
    is_full_content: bool
    metadata: Optional[dict] = None
    
    @property
    def char_count(self) -> int:
        """Return the number of characters in the extracted text."""
        return len(self.text)


@runtime_checkable
class Extractor(Protocol):
    """Protocol for file content extractors."""
    
    def can_extract(self, path: Path) -> bool:
        """Return True if this extractor can handle the given file."""
        ...
    
    def extract(self, path: Path, max_chars: int) -> ExtractedContent:
        """
        Extract text content from the file.
        
        Args:
            path: Path to the file
            max_chars: Maximum characters to extract
            
        Returns:
            ExtractedContent with the extracted text and metadata
        """
        ...


class PlainTextExtractor:
    """Extractor for plain text files."""
    
    def can_extract(self, path: Path) -> bool:
        """Plain text extractor is the fallback for all non-specialized files."""
        return True
    
    def extract(self, path: Path, max_chars: int) -> ExtractedContent:
        """Extract content from a plain text file."""
        try:
            raw = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ExtractedContent(
                text="",
                method="text",
                truncated=False,
                is_full_content=False,
            )
        
        if not isinstance(max_chars, int) or max_chars <= 0:
            return ExtractedContent(
                text=raw,
                method="text",
                truncated=False,
                is_full_content=True,
            )
        
        if len(raw) <= max_chars:
            return ExtractedContent(
                text=raw,
                method="text",
                truncated=False,
                is_full_content=True,
            )
        
        return ExtractedContent(
            text=raw[:max_chars],
            method="text",
            truncated=True,
            is_full_content=False,
        )


class PDFExtractor:
    """Extractor for PDF files."""
    
    EXTENSIONS = {".pdf"}
    
    def can_extract(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS
    
    def extract(self, path: Path, max_chars: int) -> ExtractedContent:
        """Extract text content from a PDF file."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(path))
            parts: list[str] = []
            total = 0
            truncated = False
            
            for page in reader.pages:
                t = page.extract_text() or ""
                if not t:
                    continue
                if isinstance(max_chars, int) and max_chars > 0 and total + len(t) > max_chars:
                    remaining = max_chars - total
                    if remaining > 0:
                        parts.append(t[:remaining])
                    truncated = True
                    break
                parts.append(t)
                total += len(t)
            
            text = "\n".join(parts).strip()
            return ExtractedContent(
                text=text,
                method="pdf-text",
                truncated=truncated,
                is_full_content=not truncated,
            )
        except Exception:
            return ExtractedContent(
                text="",
                method="pdf-text",
                truncated=False,
                is_full_content=False,
            )


class DocxExtractor:
    """Extractor for Microsoft Word documents."""
    
    EXTENSIONS = {".docx"}
    
    def can_extract(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS
    
    def extract(self, path: Path, max_chars: int) -> ExtractedContent:
        """Extract text content from a DOCX file."""
        text, truncated = self._extract_with_python_docx(path, max_chars)
        if text:
            return ExtractedContent(
                text=text,
                method="docx-text",
                truncated=truncated,
                is_full_content=not truncated,
            )
        
        # Fallback: parse underlying OOXML
        text, truncated = self._extract_with_lxml_fallback(path, max_chars)
        return ExtractedContent(
            text=text,
            method="docx-text",
            truncated=truncated,
            is_full_content=not truncated if text else False,
        )
    
    def _extract_with_python_docx(self, path: Path, max_chars: int) -> Tuple[str, bool]:
        """Primary extraction using python-docx."""
        try:
            import docx
            
            d = docx.Document(str(path))
            parts: list[str] = []
            used = 0
            truncated = False
            
            def append_line(line: str) -> Tuple[int, bool]:
                nonlocal used, truncated
                if not line:
                    return used, truncated
                if max_chars > 0 and used + len(line) > max_chars:
                    remaining = max_chars - used
                    if remaining > 0:
                        parts.append(line[:remaining])
                    return max_chars, True
                parts.append(line)
                return used + len(line), truncated
            
            # Body paragraphs
            for para in d.paragraphs:
                t = (para.text or "").strip()
                if not t:
                    continue
                used, truncated = append_line(t + "\n")
                if truncated:
                    break
            
            # Tables
            if not truncated:
                for table in d.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                t = (para.text or "").strip()
                                if not t:
                                    continue
                                used, truncated = append_line(t + "\n")
                                if truncated:
                                    break
                            if truncated:
                                break
                        if truncated:
                            break
                    if truncated:
                        break
            
            # Headers/footers
            if not truncated:
                for section in d.sections:
                    for hf in (section.header, section.footer):
                        for para in hf.paragraphs:
                            t = (para.text or "").strip()
                            if not t:
                                continue
                            used, truncated = append_line(t + "\n")
                            if truncated:
                                break
                        if truncated:
                            break
                    if truncated:
                        break
            
            return "".join(parts).strip(), truncated
        except Exception:
            return "", False
    
    def _extract_with_lxml_fallback(self, path: Path, max_chars: int) -> Tuple[str, bool]:
        """Fallback extraction using lxml to parse OOXML directly."""
        try:
            import zipfile
            from lxml import etree
            
            with zipfile.ZipFile(path) as zf:
                xml = zf.read("word/document.xml")
            
            root = etree.fromstring(xml)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            texts = root.xpath(".//w:t/text()", namespaces=ns)
            combined = " ".join([t.strip() for t in texts if t and t.strip()])
            combined = re.sub(r"\s+", " ", combined).strip()
            
            if not combined:
                return "", False
            
            if max_chars > 0 and len(combined) > max_chars:
                return combined[:max_chars], True
            
            return combined, False
        except Exception:
            return "", False


class XlsxExtractor:
    """Extractor for Microsoft Excel spreadsheets."""
    
    EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
    
    def can_extract(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS
    
    def extract(self, path: Path, max_chars: int) -> ExtractedContent:
        """Extract text content from an Excel file."""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(filename=str(path), read_only=True, data_only=True)
            parts: list[str] = []
            used = 0
            truncated = False
            
            for ws in wb.worksheets:
                header = f"# Sheet: {ws.title}\n"
                if max_chars > 0 and used + len(header) > max_chars:
                    truncated = True
                    break
                parts.append(header)
                used += len(header)
                
                for row in ws.iter_rows(values_only=True):
                    cells: list[str] = []
                    for v in row:
                        if v is None:
                            cells.append("")
                        else:
                            s = str(v).replace("\t", " ").replace("\r", " ").replace("\n", " ").strip()
                            cells.append(s)
                    
                    if not any(cells):
                        continue
                    
                    line = "\t".join(cells).rstrip() + "\n"
                    if max_chars > 0 and used + len(line) > max_chars:
                        remaining = max_chars - used
                        if remaining > 0:
                            parts.append(line[:remaining])
                        truncated = True
                        break
                    parts.append(line)
                    used += len(line)
                
                parts.append("\n")
                used += 1
                if truncated:
                    break
            
            text = "".join(parts).strip()
            return ExtractedContent(
                text=text,
                method="xlsx-cells",
                truncated=truncated,
                is_full_content=not truncated,
            )
        except Exception:
            return ExtractedContent(
                text="",
                method="xlsx-cells",
                truncated=False,
                is_full_content=False,
            )


class ImageExtractor:
    """Extractor for image files (returns empty content with image metadata)."""
    
    EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".tiff", ".bmp", ".heic"}
    
    def can_extract(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS
    
    def extract(self, path: Path, max_chars: int) -> ExtractedContent:
        """Images return empty text (could be extended for OCR in the future)."""
        return ExtractedContent(
            text="",
            method="image",
            truncated=False,
            is_full_content=False,
        )


class GoogleStubExtractor:
    """Extractor for Google Workspace stub files (.gdoc, .gsheet, .gslides)."""
    
    EXTENSIONS = {".gdoc", ".gsheet", ".gslides", ".gform"}
    _GOOGLE_ID_RE = re.compile(r"/d/([a-zA-Z0-9_-]+)")
    _GOOGLE_ID_QS_RE = re.compile(r"[?&]id=([a-zA-Z0-9_-]+)")
    
    def can_extract(self, path: Path) -> bool:
        return path.suffix.lower() in self.EXTENSIONS
    
    def extract(self, path: Path, max_chars: int) -> ExtractedContent:
        """
        Extract the Google file ID from a stub file.
        
        Stub files are JSON pointers to Google Drive files. The actual content
        would need to be fetched via Drive API (handled separately in drive.py).
        """
        google_id = self._extract_google_id(path)
        return ExtractedContent(
            text="",
            method="google-stub",
            truncated=False,
            is_full_content=False,
            metadata={"google_id": google_id} if google_id else None,
        )
    
    def _extract_google_id(self, path: Path) -> Optional[str]:
        """Extract the Google Drive file ID from a stub file."""
        try:
            raw = path.read_text(encoding="utf-8", errors="ignore")
            obj = json.loads(raw or "{}")
        except Exception:
            raw, obj = None, {}
        
        # Check common JSON keys
        for k in ("id", "doc_id", "resource_id"):
            v = obj.get(k)
            if isinstance(v, str) and v.strip():
                return v.strip()
        
        # Check URL fields
        for k in ("url", "open_url", "alternate_link", "alternateLink", "app_url"):
            v = obj.get(k)
            if isinstance(v, str) and v:
                m = self._GOOGLE_ID_RE.search(v) or self._GOOGLE_ID_QS_RE.search(v)
                if m:
                    return m.group(1)
        
        # Last resort: search raw content
        if raw:
            m = self._GOOGLE_ID_RE.search(raw) or self._GOOGLE_ID_QS_RE.search(raw)
            if m:
                return m.group(1)
        
        return None


# Registry of extractors in priority order (most specific first)
_EXTRACTORS: list[Extractor] = [
    PDFExtractor(),
    DocxExtractor(),
    XlsxExtractor(),
    ImageExtractor(),
    GoogleStubExtractor(),
    PlainTextExtractor(),  # Fallback, must be last
]


def get_extractor(path: Path) -> Extractor:
    """
    Get the appropriate extractor for a file.
    
    Args:
        path: Path to the file
        
    Returns:
        The first matching extractor (PlainTextExtractor as fallback)
    """
    for extractor in _EXTRACTORS:
        if extractor.can_extract(path):
            return extractor
    return PlainTextExtractor()


def extract_text(path: Path, *, max_chars: int) -> ExtractedContent:
    """
    Extract text content from a file using the appropriate extractor.
    
    This is the main entry point for file extraction.
    
    Args:
        path: Path to the file
        max_chars: Maximum characters to extract
        
    Returns:
        ExtractedContent with the extracted text and metadata
    """
    extractor = get_extractor(path)
    return extractor.extract(path, max_chars)


def is_google_stub(path: Path) -> bool:
    """Return True if the file is a Google Workspace stub file."""
    return path.suffix.lower() in GoogleStubExtractor.EXTENSIONS


def extract_google_id_from_stub(path: Path) -> Optional[str]:
    """Extract the Google Drive file ID from a stub file."""
    extractor = GoogleStubExtractor()
    result = extractor.extract(path, 0)
    if result.metadata:
        return result.metadata.get("google_id")
    return None
