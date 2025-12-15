from __future__ import annotations

import json
import mimetypes
import os
import re
from pathlib import Path
from typing import Optional, Tuple

from pypdf import PdfReader


_GOOGLE_STUB_EXTS = {".gdoc", ".gsheet", ".gslides", ".gform"}
_GOOGLE_ID_RE = re.compile(r"/d/([a-zA-Z0-9_-]+)")
_GOOGLE_ID_QS_RE = re.compile(r"[?&]id=([a-zA-Z0-9_-]+)")


def is_google_stub(path: Path) -> bool:
    return path.suffix.lower() in _GOOGLE_STUB_EXTS


def extract_google_id_from_stub(path: Path) -> Optional[str]:
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        obj = json.loads(raw or "{}")
    except Exception:
        raw, obj = None, {}

    for k in ("id", "doc_id", "resource_id"):
        v = obj.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()

    for k in ("url", "open_url", "alternate_link", "alternateLink", "app_url"):
        v = obj.get(k)
        if isinstance(v, str) and v:
            m = _GOOGLE_ID_RE.search(v) or _GOOGLE_ID_QS_RE.search(v)
            if m:
                return m.group(1)

    if raw:
        m = _GOOGLE_ID_RE.search(raw) or _GOOGLE_ID_QS_RE.search(raw)
        if m:
            return m.group(1)

    return None


def sniff_mime_type(path: Path) -> Optional[str]:
    mt, _ = mimetypes.guess_type(str(path))
    return mt


def read_text_file(path: Path, max_chars: int) -> Tuple[str, bool]:
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return "", False
    if not isinstance(max_chars, int) or max_chars <= 0:
        return raw, False
    if len(raw) <= max_chars:
        return raw, False
    return raw[:max_chars], True


def extract_pdf_text(path: Path, max_chars: int) -> Tuple[str, bool]:
    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        total = 0
        truncated = False
        for page in reader.pages:
            t = page.extract_text() or ""
            if not t:
                continue
            if isinstance(max_chars, int) and max_chars > 0 and total + len(t) > max_chars:
                remaining = max_chars - total
                if remaining > 0:
                    parts.append(t[:remaining])
                truncated = True
                break
            parts.append(t)
            total += len(t)
        return "\n".join(parts).strip(), truncated
    except Exception:
        return "", False


def extract_xlsx_text(path: Path, max_chars: int) -> Tuple[str, bool]:
    """
    Extract a text preview from an XLSX/XLSM file by reading cell values.
    Returns (text, truncated).
    """
    try:
        from openpyxl import load_workbook

        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
        parts: list[str] = []
        used = 0
        truncated = False

        for ws in wb.worksheets:
            header = f"# Sheet: {ws.title}\n"
            if max_chars > 0 and used + len(header) > max_chars:
                truncated = True
                break
            parts.append(header)
            used += len(header)

            for row in ws.iter_rows(values_only=True):
                # Convert each cell to a compact, readable string.
                cells: list[str] = []
                for v in row:
                    if v is None:
                        cells.append("")
                    else:
                        s = str(v).replace("\t", " ").replace("\r", " ").replace("\n", " ").strip()
                        cells.append(s)
                # Skip entirely empty rows.
                if not any(cells):
                    continue
                line = "\t".join(cells).rstrip() + "\n"
                if max_chars > 0 and used + len(line) > max_chars:
                    remaining = max_chars - used
                    if remaining > 0:
                        parts.append(line[:remaining])
                    truncated = True
                    break
                parts.append(line)
                used += len(line)

            parts.append("\n")
            used += 1
            if truncated:
                break

        text = "".join(parts).strip()
        return text, truncated
    except Exception:
        return "", False


def extract_docx_text(path: Path, max_chars: int) -> Tuple[str, bool]:
    """
    Extract a text preview from a DOCX by reading paragraph text.
    Returns (text, truncated).
    """
    def _append_line(parts: list[str], used: int, truncated: bool, line: str) -> tuple[int, bool]:
        if not line:
            return used, truncated
        if max_chars > 0 and used + len(line) > max_chars:
            remaining = max_chars - used
            if remaining > 0:
                parts.append(line[:remaining])
            return max_chars, True
        parts.append(line)
        return used + len(line), truncated

    try:
        import docx  # python-docx

        d = docx.Document(str(path))
        parts: list[str] = []
        used = 0
        truncated = False

        # Body paragraphs
        for para in d.paragraphs:
            t = (para.text or "").strip()
            if not t:
                continue
            used, truncated = _append_line(parts, used, truncated, t + "\n")
            if truncated:
                break

        # Tables (common for templates and risk assessments)
        if not truncated:
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            t = (para.text or "").strip()
                            if not t:
                                continue
                            used, truncated = _append_line(parts, used, truncated, t + "\n")
                            if truncated:
                                break
                        if truncated:
                            break
                    if truncated:
                        break
                if truncated:
                    break

        # Headers/footers
        if not truncated:
            for section in d.sections:
                for hf in (section.header, section.footer):
                    for para in hf.paragraphs:
                        t = (para.text or "").strip()
                        if not t:
                            continue
                        used, truncated = _append_line(parts, used, truncated, t + "\n")
                        if truncated:
                            break
                    if truncated:
                        break
                    for table in hf.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    t = (para.text or "").strip()
                                    if not t:
                                        continue
                                    used, truncated = _append_line(parts, used, truncated, t + "\n")
                                    if truncated:
                                        break
                                if truncated:
                                    break
                            if truncated:
                                break
                        if truncated:
                            break
                    if truncated:
                        break
                if truncated:
                    break

        text = "".join(parts).strip()
        if text:
            return text, truncated
    except Exception:
        pass

    # Fallback: parse the underlying OOXML for any text runs (covers some documents python-docx won't expose).
    try:
        import zipfile

        from lxml import etree

        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
        root = etree.fromstring(xml)
        texts = root.xpath(".//w:t/text()", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
        combined = " ".join([t.strip() for t in texts if t and t.strip()])
        combined = re.sub(r"\s+", " ", combined).strip()
        if not combined:
            return "", False
        if max_chars > 0 and len(combined) > max_chars:
            return combined[:max_chars], True
        return combined, False
    except Exception:
        return "", False


def sanitize_folder_name(name: str) -> str:
    name = (name or "").strip()
    name = re.sub(r"\s+", " ", name)
    name = name.strip(" .")
    if not name:
        raise ValueError("Folder name is empty.")
    if any(sep in name for sep in ("/", "\\", "\0")):
        raise ValueError(f"Invalid folder name: {name!r}")
    return name


def managed_index_update(existing: str, managed_markdown: str) -> str:
    start = "<!-- SMARTSORTER MANAGED SECTION START -->"
    end = "<!-- SMARTSORTER MANAGED SECTION END -->"
    block = f"{start}\n{managed_markdown.rstrip()}\n{end}\n"

    if start in existing and end in existing and existing.index(start) < existing.index(end):
        pre = existing.split(start, 1)[0].rstrip() + "\n\n"
        post = existing.split(end, 1)[1].lstrip()
        return pre + block + ("\n" + post if post else "")

    existing = existing.rstrip()
    if existing:
        existing += "\n\n"
    return existing + block


def env_truthy(name: str) -> bool:
    v = os.environ.get(name, "").strip().lower()
    return v in {"1", "true", "yes", "y", "on"}


def is_ignorable_file_name(name: str) -> bool:
    """
    Return True for common OS/system metadata files that should not be processed.
    """
    n = (name or "").strip()
    if not n:
        return True
    if n in {".DS_Store", "Thumbs.db", "desktop.ini"}:
        return True
    # AppleDouble "resource fork" files created on non-HFS filesystems and some sync tools.
    if n.startswith("._"):
        return True
    # Generic hidden dotfiles (metadata) are skipped by default.
    if n.startswith("."):
        return True
    return False
